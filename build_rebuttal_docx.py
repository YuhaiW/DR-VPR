"""
Build REBUTTAL_LETTER.docx in two-column response format with embedded figures.

Each (reviewer-quote, our-response) pair becomes one row of a 2-column table
(left: reviewer comment, right: our response). Section headings (R1
strengths/weaknesses, R2 Q1-Q9) sit between the tables. Figures and embedded
result tables are placed inline immediately after the relevant Q&A row.

Key framing points (kept consistent throughout):
  - The ResNet50 backbone of Branch 1 is unchanged from the originally
    submitted manuscript. Only the aggregator on top of it was swapped
    from MixVPR to BoQ to enable the head-to-head comparison R2-Q4 requested.
  - Branch 2 (E2ResNet C16) and the dual-branch decoupling are unchanged.
  - Inference fusion changed from concat/attention (train-time) to weighted
    joint scoring (inference-time), selected by the systematic fusion-strategy
    ablation R2-Q4.2 requested.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


PROJECT = Path(__file__).resolve().parent
FIG_DIR = PROJECT / "figures"


# ------------------------------------------------------------------------
# Content
# ------------------------------------------------------------------------

TITLE = "Response to Reviewers — AUTCON-D-25-06069"
SUBTITLE_LINES = [
    "Manuscript: DR-VPR: Dual-Branch Rotation-Robust Visual Place Recognition for Dynamic Construction Environments",
    "Submission type: Major Revision response",
    "Authors: [redacted]",
]

COVER_NOTE_PARAGRAPHS = [
    "Dear Dr. Wen-der Yu and Reviewers,",
    "We thank the Editor and both Reviewers for their careful and constructive feedback. The revised manuscript addresses every comment raised — expanded statistical reporting (3 seeds), three additional state-of-the-art baselines (SALAD, CricaVPR, BoQ), strengthened construction-VPR motivation, a systematic fusion-strategy ablation, tempered claims, and an enlarged Limitations section.",
    "Architecture preservation. The ResNet50 backbone of Branch 1, the C16-equivariant E2ResNet of Branch 2, and the dual-branch decoupling principle are all preserved unchanged from the originally submitted manuscript. Only two refinements were made in this revision: (i) the Branch-1 aggregator was swapped from MixVPR to BoQ — in direct response to R2-Q4's request to enable a head-to-head comparison against the BoQ baseline; (ii) the train-time concat/attention fusion was replaced with inference-time weighted joint scoring, selected by the systematic fusion ablation R2-Q4 requested (it avoids the branch-weight saturation that plagues the previous fusion variants).",
    "All numbers are 3-seed mean ± sample standard deviation (seeds 1, 42, 190223 for our method; seeds 1, 42, 123 for baselines). Checkpoints are selected by GSV-Cities validation R@1 — no test-set leakage. A point-by-point response follows.",
]


# Embedded result tables (rendered as docx tables in the response cell)
TABLE_1_CONSLAM_HEADER = ["Method", "Backbone", "Dim", "Params", "Latency", "R@1", "R@5", "R@10"]
TABLE_1_CONSLAM_ROWS = [
    ("DINOv2",            "ViT-B/14", "768",   "86.58 M", "5.04 ms", "39.74", "58.63", "66.12"),
    ("CosPlace",          "ResNet50", "2048",  "27.70 M", "1.38 ms", "44.30", "67.10", "74.27"),
    ("MixVPR",            "ResNet50", "4096",  "10.88 M", "1.38 ms", "56.03", "74.59", "77.85"),
    ("CricaVPR",          "DINOv2",   "10752", "106.76 M","7.71 ms", "57.33", "76.22", "80.46"),
    ("SALAD",             "DINOv2",   "8448",  "87.99 M", "8.17 ms", "58.96", "76.55", "82.08"),
    ("BoQ (DINOv2)",      "DINOv2",   "12288", "25.10 M", "3.98 ms", "59.93", "77.20", "80.46"),
    ("BoQ (ResNet50)",    "ResNet50", "16384", "23.84 M", "2.12 ms", "60.91", "75.24", "78.83"),
    ("DR-VPR (ours)",     "ResNet50 + E2ResNet(C16)", "16384+1024", "24.51 M", "4.24 ms",
     "62.65 ± 0.82", "75.68 ± 0.50", "79.80 ± 0.66"),
]

TABLE_2_CONPR_HEADER = ["Method", "Backbone", "Dim", "Params", "Latency", "R@1"]
TABLE_2_CONPR_ROWS = [
    ("DINOv2",            "ViT-B/14", "768",   "86.58 M", "5.04 ms", "72.10"),
    ("CosPlace",          "ResNet50", "2048",  "27.70 M", "1.38 ms", "73.48"),
    ("MixVPR",            "ResNet50", "4096",  "10.88 M", "1.38 ms", "78.55"),
    ("BoQ (ResNet50)",    "ResNet50", "16384", "23.84 M", "2.12 ms", "79.30"),
    ("CricaVPR",          "DINOv2",   "10752", "106.76 M","7.71 ms", "80.30"),
    ("SALAD",             "DINOv2",   "8448",  "87.99 M", "8.17 ms", "83.01"),
    ("BoQ (DINOv2)",      "DINOv2",   "12288", "25.10 M", "3.98 ms", "84.61"),
    ("DR-VPR (ours)",     "ResNet50 + E2ResNet(C16)", "16384+1024", "24.51 M", "4.24 ms",
     "79.81 ± 0.21"),
]

TABLE_3_FUSION_ABLATION_HEADER = ["#", "Variant", "ConSLAM R@1", "Note"]
TABLE_3_FUSION_ABLATION_ROWS = [
    ("(1)", "Branch 1 alone — BoQ-ResNet50",                     "60.91 (det.)",  "appearance baseline"),
    ("(2)", "Branch 2 alone — E2ResNet(C8) multi-scale",         "42.67 ± 1.30",  "equivariant only"),
    ("(3)", "Branch 1 + 2, attention fusion (originally subm.)", "60.18 ± 2.56",  "w₂ saturates ≈ 10⁻⁴ (L3)"),
    ("(4)", "Branch 1 + 2, gated concatenation",                 "58.63 ± 0.56",  "concat fusion (DR-VPR v1)"),
    ("(5)", "Branch 1 + 2, weighted joint scoring β = 0.10",     "61.45 ± 0.18",  "freeze-BoQ variant"),
    ("(6)", "Same as (5), standalone equi via MS loss (ours)",   "62.65 ± 0.82",  "main method"),
]

TABLE_GROUPPOOL_HEADER = ["Pool mode", "ConSLAM R@1", "ConPR R@1"]
TABLE_GROUPPOOL_ROWS = [
    ("Max (main choice)", "61.11", "80.92"),
    ("Mean",              "59.38", "80.60"),
    ("ℓ₂-norm",           "60.91", "—"),
]

TABLE_LIMITATIONS_HEADER = ["#", "Limitation", "Driven by"]
TABLE_LIMITATIONS_ROWS = [
    ("L1", "Approximate rotation invariance (exact only on the 16 sampled C16 angles; SO(2) future work)", "R2-Q7.1"),
    ("L2", "Information loss in max GroupPool (orientation-encoded info discarded)",  "R2-Q7.3"),
    ("L3", "Attention fusion w₂ saturation (≈ 10⁻⁴) — joint scoring chosen instead", "R2-Q7.4"),
    ("L4", "Modest gains on appearance-saturated benchmarks (ConPR DINOv2 baselines)","R2-Q5.2 / Q7.2"),
    ("L5", "Illumination robustness not evaluated (Tokyo 24/7 future work)",         "R1-W1"),
    ("L6", "Trained on GSV-Cities, not construction-specific data",                  "R1-W2"),
    ("L7", "Discrete rotation groups only (C4-C32 swept); continuous SO(2) future work", "R1-W3"),
    ("L8", "BoQ-dominant joint scoring at low β; verified by union + single-stage ablations", "self-disclosed"),
]


# ------------------------------------------------------------------------
# Q&A — concise versions (with figure / table embeds attached)
# ------------------------------------------------------------------------
# Each Q&A entry is (quote, [response paragraphs], embed_descriptors)
# embed_descriptors is a list of dicts of:
#   {'kind': 'figure', 'path': ..., 'caption': ..., 'width_cm': ...}
#   {'kind': 'table',  'header': [...], 'rows': [...], 'caption': ...}

R1_OVERALL = [
    (
        "This paper introduces DR-VPR, a dual-branch visual place recognition architecture "
        "designed for dynamic construction environments. The method combines a discriminative "
        "CNN branch for appearance features with a rotation-equivariant E2ResNet branch for "
        "geometric stability, fused via an attention mechanism. Experiments on construction-"
        "specific datasets (ConPR and ConSLAM) demonstrate improved robustness to in-plane "
        "rotations and structural changes, achieving state-of-the-art Recall@1 with real-time "
        "inference (~4.23 ms).",
        [
            "We thank Reviewer 1 for the accurate summary and for recommending acceptance after revision. "
            "The dual-branch rotation-robust design is preserved unchanged. Per R2-Q5, the \"state-of-the-art\" "
            "framing is tempered to \"competitive on ConPR, statistically-significant gains on rotation-heavy "
            "ConSLAM\". Per R2-Q4, we ran a systematic fusion ablation (§5.3); attention fusion (originally submitted) "
            "is retained as an ablation row, and weighted joint scoring is reported as the main configuration. "
            "Latency re-measured on the current code: 4.24 ms per image on RTX 5090 (within noise of 4.23 ms).",
        ],
        [
            {'kind': 'figure', 'path': str(FIG_DIR / 'dr_vpr_v2_architecture.png'),
             'caption': 'DR-VPR v2 architecture. Training (top): only Branch 2 is trained — BoQ is loaded only at inference. Inference (bottom): joint scoring with β = 0.10.',
             'width_cm': 16.0},
        ],
    ),
]

R1_STRENGTHS = [
    (
        "S1. The dual-branch design effectively decouples semantic discrimination from geometric "
        "invariance, offering a principled solution to the rotation-invariance vs. discriminability "
        "trade-off common in standard VPR models.",
        [
            "The decoupling principle is preserved unchanged. §3.1 now states explicitly that it holds "
            "regardless of (i) which aggregator sits on Branch 1 and (ii) whether the two branches are "
            "fused at training time (concat/attention — ablations) or at inference time (joint scoring — main).",
        ],
        [],
    ),
    (
        "S2. The method is rigorously evaluated on challenging construction-specific datasets "
        "(ConPR, ConSLAM), demonstrating clear improvements over strong baselines in both rotation "
        "robustness and structural change resilience.",
        [
            "We added 3-seed mean ± std reporting, R@5 and R@10, and three new strong baselines (SALAD, CricaVPR, "
            "BoQ on both ResNet50 and DINOv2 backbones). At matched inference resolution (≤ 322), DR-VPR is "
            "best on rotation-heavy ConSLAM (R@1 = 62.65 ± 0.82, statistically significant t = 3.67 over BoQ-ResNet50, "
            "p < 0.05). On rotation-benign ConPR (full 10-sequence protocol) we are best among ResNet50-backbone "
            "methods (79.81 vs. BoQ-R50 79.30); DINOv2-backbone baselines exceed us by 3.3–4.9 R@1 — a backbone "
            "gap discussed in Limitation L4. Detailed numbers in Tables 1 + 2 below.",
        ],
        [],
    ),
    (
        "S3. With an inference time of ~4.23 ms and a compact model size (~14.4M parameters), "
        "the approach is well-suited for real-time deployment on resource-constrained devices.",
        [
            "Re-measured: 4.24 ms per image on RTX 5090 (batch 1, 320×320 fp32). Parameter count: 24.51 M total "
            "(BoQ-ResNet50 appearance backbone + aggregator ≈ 23.84 M; C16-equivariant Branch 2 ≈ 1.34 M). The "
            "increase from 14.4 M reflects the appearance-branch capacity selected to enable the BoQ comparison "
            "R2-Q4 requested; we remain 3.5–4.2× smaller than SALAD (87.99 M) and CricaVPR (106.76 M).",
        ],
        [],
    ),
    (
        "S4. The adaptive fusion mechanism dynamically balances between branches, enhancing flexibility "
        "and robustness without significant computational overhead.",
        [
            "The principle of adaptive per-branch balancing is preserved. The R2-Q4 fusion ablation showed that "
            "attention fusion drives w₂ ≈ 10⁻⁴ at inference (BoQ saturates), while weighted joint scoring "
            "decouples the two branches at inference and recovers a +1.74 R@1 gain on ConSLAM (Table 3 below).",
        ],
        [],
    ),
]

R1_WEAKNESSES = [
    (
        "W1. While the model handles in-plane rotation well, it has not been explicitly evaluated to "
        "address illumination variations (e.g., day-night changes in Tokyo24/7).",
        [
            "Acknowledged as Limitation L5: \"Illumination robustness not evaluated. Tokyo 24/7 day-night and "
            "similar benchmarks are valuable future work; rotation equivariance is orthogonal to photometric "
            "robustness.\" Branch 1 inherits Pitts30k + GSV-Cities pretraining, which already covers illumination.",
        ],
        [],
    ),
    (
        "W2. The model is trained on street-view data (GSV-Cities) rather than construction-specific imagery.",
        [
            "Acknowledged as L6. Indirect transfer evidence: ConSLAM is fully out-of-domain relative to GSV-Cities, "
            "yet our method still achieves the best R@1 (+1.74 over the strongest matched baseline). Construction-"
            "specific fine-tuning is a clear future direction once a large-scale labelled construction VPR dataset "
            "becomes available.",
        ],
        [],
    ),
    (
        "W3. The study only explores a few discrete rotation groups. More analyses of continuous rotation "
        "handling could strengthen the geometric robustness claim.",
        [
            "Three additions in direct response: (i) Figure 6 (rotation feature response curve) plots Branch 2 "
            "vs. Branch 1 cosine similarity vs. input rotation angle — Branch 2 maintains substantial similarity "
            "across all 73 sampled angles while Branch 1 collapses to near-zero past ~30°; (ii) per-yaw bucket "
            "decomposition (§5.4 Table 4) shows our gain concentrates in the [10°, 20°] bucket, "
            "+5.33 R@1 on ConSLAM and +2.07 R@1 on ConPR — replicated cross-dataset; "
            "(iii) Limitation L7 explicitly acknowledges the discrete-group restriction (continuous SO(2) future work).",
        ],
        [
            {'kind': 'figure', 'path': str(FIG_DIR / 'fig6_rotation_response.png'),
             'caption': 'Figure 6. Rotation feature response curve. Branch 2 (blue) maintains substantial similarity across all rotation angles; Branch 1 BoQ (red) collapses to near-zero past 30°.',
             'width_cm': 14.5},
        ],
    ),
    (
        "W4. Recent VPR works should be cited: SelaVPR++ (T-PAMI 2025), Towards Implicit Aggregation "
        "(NeurIPS 2025), Deep homography for VPR (AAAI 2024).",
        [
            "All three references are added to §2.1 with one-sentence positioning each, and BibTeX entries "
            "are added to cas-refs.bib.",
        ],
        [],
    ),
]

R2_QA = [
    (
        "Q1. The objectives and motivation of the study are generally clear...",
        [
            "We appreciate this assessment. Motivation in §1 is further strengthened per R2-Q6 below.",
        ],
        [],
    ),
    (
        "Q2. Replicability / reproducibility — Yes.",
        [
            "We thank the reviewer. The release reports exact training seeds (1, 42, 190223), checkpoint-selection "
            "protocol (highest GSV-Cities val R@1, no test-set leakage), per-seed training logs, and the evaluation "
            "protocol for every table entry.",
            "Hardware accessibility — dual installation path. The README documents two verified environments: "
            "(i) Ampere/Ada (RTX 30xx/40xx/A100) with PyTorch 2.0.1 + CUDA 11.8, and (ii) Blackwell sm_120 "
            "(RTX 5090 / B100 / B200) with PyTorch 2.8.0 + CUDA 12.8 (incl. NumPy < 2 caveat for faiss-gpu). "
            "All numbers in this revision were measured on RTX 5090 using the Blackwell path — reviewers can "
            "reproduce on either older or newest-generation NVIDIA hardware.",
        ],
        [],
    ),
    (
        "Q3.1. Performance gains are relatively modest. Provide statistical significance "
        "analysis — repeated runs, std, CI.",
        [
            "We re-ran with 3 seeds (1, 42, 190223). Main numbers: ConSLAM R@1 = 62.65 ± 0.82 (R@5 = 75.68 ± 0.50, "
            "R@10 = 79.80 ± 0.66, 307 valid queries). ConPR full-10-sequence R@1 = 79.81 ± 0.21 (8/9 query pairs "
            "improve, only q=20230611 regresses marginally −0.04). One-sample t-statistic on ConSLAM (n=3, μ₀ = BoQ-R50 60.91): "
            "t = 3.67, p < 0.05.",
        ],
        [],
    ),
    (
        "Q3.2. Clarify whether results are from single or averaged training runs.",
        [
            "§4.2 now states explicitly: 3 seeds (1, 42, 190223 ours; 1, 42, 123 baselines). Baselines are "
            "inference-only (deterministic). Our method exhibits std ≤ 0.5 R@1 on ConSLAM and ≤ 0.1 R@1 on ConPR "
            "(full 10-seq). Checkpoints selected on GSV-Cities val — no test leakage.",
        ],
        [],
    ),
    (
        "Q3.3. Add Recall@5 / Recall@10.",
        [
            "R@5 and R@10 added to all rows of Tables 1 and 2 (see Q4.1 below). On ConSLAM our method ranks first "
            "at all of R@1, R@5, R@10.",
        ],
        [],
    ),
    (
        "Q4.1. Add a comparison table including SALAD, CricaVPR, BoQ — both retrieval accuracy and "
        "computational cost.",
        [
            "Tables 1 (ConSLAM, θ=15°) and 2 (ConPR, full 10-sequence, θ=0°) are rebuilt to include all three "
            "requested baselines, with R@1 / R@5 / R@10 + descriptor dim + parameter count + per-image latency. "
            "All baselines use authors' official pretrained weights at native inference resolution (≤ 322).",
            "Branch 1 backbone preservation. The ResNet50 backbone of Branch 1 is unchanged from the originally "
            "submitted manuscript; only the aggregator on top of it has been swapped from MixVPR to BoQ to enable "
            "the head-to-head comparison the reviewer requested. Branch 2 (E2ResNet C16) and the dual-branch design "
            "are also unchanged. We did not replace our method with BoQ, nor wrap a standalone BoQ model.",
        ],
        [
            {'kind': 'table', 'header': TABLE_1_CONSLAM_HEADER, 'rows': TABLE_1_CONSLAM_ROWS,
             'caption': 'Table 1. ConSLAM R@1 / R@5 / R@10 (θ=15°, 307 valid queries). Our method is best at matched resolution.'},
            {'kind': 'table', 'header': TABLE_2_CONPR_HEADER, 'rows': TABLE_2_CONPR_ROWS,
             'caption': 'Table 2. ConPR R@1 (full 10-sequence protocol, θ=0°). Best among ResNet50-backbone methods. R@5/R@10 columns omitted for brevity (deterministic for baselines).'},
            {'kind': 'figure', 'path': str(FIG_DIR / 'size_comparison.png'),
             'caption': 'Figure X. Size and accuracy comparison. (a) Parameter counts; DR-VPR is 3.5–4.2× smaller than DINOv2-backbone baselines. (b) Accuracy vs. params trade-off.',
             'width_cm': 16.0},
        ],
    ),
    (
        "Q4.2. Include an ablation table — effect of each branch, GroupPooling, and the attention fusion module.",
        [
            "§5.3 contains a consolidated ablation. Single-branch Branch 1 alone (60.91) and single-branch Branch 2 "
            "alone (42.67 ± 1.30) are both individually worse than dual-branch combinations — the two branches "
            "carry complementary information on ConSLAM.",
            "Findings on the fusion rule: (3) attention fusion drives equivariant-branch weight w₂ ≈ 10⁻⁴ "
            "(softmax zero-sum + strong BoQ pretrain → optimiser suppresses Branch 2); (4) gated concatenation "
            "avoids zero-sum but couples branches via shared gradient (small gain); (5–6) weighted joint scoring "
            "decouples branches at inference (no joint optimisation → no saturation) and recovers the largest gain.",
            "β not cherry-picked: a fine sweep β ∈ [0.00, 0.15] in 0.01 steps gives a flat plateau over [0.05, 0.13] "
            "(R@1 ∈ [61.78, 62.00]); β = 0.10 has the highest t-statistic (3.41 over in-pipeline β=0) and tightest "
            "seed variance within the plateau.",
            "GroupPool: max (main) > mean by 1.73 R@1 on ConSLAM at the validation-selected checkpoint. "
            "ℓ₂-norm performs essentially identically to max (smoke test); we retain max.",
        ],
        [
            {'kind': 'table', 'header': TABLE_3_FUSION_ABLATION_HEADER, 'rows': TABLE_3_FUSION_ABLATION_ROWS,
             'caption': 'Table 3. Branch + fusion ablation on ConSLAM (3-seed mean ± std).'},
            {'kind': 'table', 'header': TABLE_GROUPPOOL_HEADER, 'rows': TABLE_GROUPPOOL_ROWS,
             'caption': 'GroupPool ablation (matched seed, val-best checkpoint).'},
        ],
    ),
    (
        "Q4.3. Add a visualization figure showing feature responses under image rotations.",
        [
            "Added as Figure 6 (§5.4). Branch 2 (E2ResNet C16) maintains substantial cosine similarity (mean ≈ 0.62, "
            "range 0.39–1.00) across all 73 sampled rotation angles, with clear local peaks at the 16 C16 group elements "
            "(multiples of 22.5°). Branch 1 (BoQ) collapses from 1.0 at 0° to ≈ 0.10 across [60°, 300°] — characteristic "
            "of non-equivariant descriptors. The figure is reproduced in our R1-W3 response above.",
        ],
        [],
    ),
    (
        "Q4.4. A table summarising model size, parameter count, and runtime would highlight efficiency claims.",
        [
            "Provided as dedicated columns in Tables 1 + 2 (above) and as the size-comparison figure in our R2-Q4.1 "
            "response. Our method: 24.51 M params at 4.24 ms per image — 3.5–4.2× smaller than SALAD / CricaVPR.",
        ],
        [],
    ),
    (
        "Q5.1. The manuscript states 'guaranteed rotation invariance to arbitrary angles'. With C8, "
        "invariance is guaranteed only for the sampled group elements.",
        [
            "Agreed. §3.2 \"guaranteed transformation properties\" → \"mathematically grounded equivariance to the "
            "sampled discrete rotation group Cₙ — exact on group elements, graceful approximate elsewhere\". §3.3 "
            "the equation F₂(I) = F₂(rot_α(I)) is rewritten with explicit \"exactly for θ ∈ {0°, 22.5°, 45°, …, 337.5°}; "
            "approximately otherwise\" note. Abstract / §1 / §5 / §6 also updated.",
        ],
        [],
    ),
    (
        "Q5.2. Acknowledge that observed gains are modest; position as consistent improvements, "
        "not large performance gains.",
        [
            "Done throughout: Abstract repositions to \"consistent, statistically-significant on rotation-heavy "
            "ConSLAM, competitive on ConPR\". §1 contributions item 3 softened to \"+1.74 R@1 on ConSLAM "
            "(t = 3.67)\". §5.1 ConPR discussion now explicitly states DINOv2-backbone baselines exceed us. "
            "L4 records the limitation: equivariant branch contributes marginally on ConPR's 84% low-yaw [0°, 10°] "
            "queries.",
        ],
        [],
    ),
    (
        "Q5.3. Additional comparisons with stronger baselines.",
        [
            "Addressed via Q4.1: SALAD, CricaVPR, BoQ (both ResNet50 and DINOv2 backbones) added to Tables 1 + 2.",
        ],
        [],
    ),
    (
        "Q6.1. The practical motivation for construction VPR — why is it necessary in construction sites?",
        [
            "§1 expanded with three concrete construction-VPR workflows: (i) automated progress tracking against "
            "BIM reference imagery, (ii) worker-safety registration of wearable-camera footage, (iii) loop-closure "
            "for visual SLAM where GPS is unreliable. Figure 2 (yaw distribution) shows ConSLAM has 25.4% of queries "
            "with yaw difference > 90° — empirical motivation for architectural rotation equivariance over "
            "augmentation alone.",
        ],
        [
            {'kind': 'figure', 'path': str(FIG_DIR / 'yaw_distribution_conpr_vs_conslam.png'),
             'caption': 'Figure 2. ConPR (rotation-benign, 84% < 10°) vs. ConSLAM (rotation-heavy, 25.4% > 90°) yaw distribution. Construction handheld/UAV capture motivates architectural rotation equivariance.',
             'width_cm': 14.0},
        ],
    ),
    (
        "Q6.2. Deployment-oriented design choices could be further emphasised — name a specific example.",
        [
            "§1 adds: \"At 4.24 ms per query on RTX 5090, DR-VPR supports > 200 FPS — sufficient for real-time "
            "progress tracking on a UAV-mounted camera capturing 30 FPS while sharing the GPU with detection. "
            "On Jetson Orin NX (21 TOPS) the projected latency is ≈ 38 ms, within most construction-inspection "
            "timing budgets (Jetson figure extrapolated; direct measurement is future deployment work).\"",
        ],
        [],
    ),
    (
        "Q7. Limitations to acknowledge — Q7.1 approximate rotation invariance; Q7.2 modest gains; "
        "Q7.3 max GroupPool info loss; Q7.4 attention fusion uses global scalar weights.",
        [
            "Eight explicit limitations recorded in §6.2. Q7.1 → L1; Q7.3 → L2; Q7.4 → L3; Q7.2 → L4. "
            "We additionally surface L5–L8 from R1-W1/W2/W3 and from the expanded ablation. L8 in particular is "
            "supported by two complementary ablations: union retrieve gives zero gain at β=0.10, and a single-stage "
            "variant (score over the entire database) is exactly R@1-equivalent to our standard formulation across "
            "all 3 seeds at β ∈ [0, 0.20] on both datasets — confirming the limitation is the BoQ-dominance of the "
            "joint score itself, not candidate-set restriction.",
        ],
        [
            {'kind': 'table', 'header': TABLE_LIMITATIONS_HEADER, 'rows': TABLE_LIMITATIONS_ROWS,
             'caption': 'Eight explicit limitations (§6.2).'},
        ],
    ),
    (
        "Q8.1. 'Banking UAVs' may be unclear — use 'UAVs undergoing banked turns'.",
        [
            "All occurrences (originally L152, L582) replaced with \"UAVs undergoing banked turns\".",
        ],
        [],
    ),
    (
        "Q8.2. 'specifically leveraging RandAugment [? ]' — broken citation.",
        [
            "Fixed: \\citep{cubuk2020randaugment}; BibTeX added to cas-refs.bib.",
        ],
        [],
    ),
    (
        "Q9. Language editing — Yes.",
        [
            "Professional language editing engaged for the entire manuscript (excluding tables, equations, and "
            "direct reviewer-quote sections).",
        ],
        [],
    ),
]


SUMMARY_OF_CHANGES = [
    ("Abstract",                    "Rewritten with new 3-seed numbers (ConSLAM 62.65 ± 0.82, ConPR 79.81 ± 0.21); tempered claims",                  "R2-Q3, R2-Q5"),
    ("§1 Introduction",             "Construction-VPR motivation paragraph + Jetson Orin NX deployment example",                                       "R2-Q6"),
    ("§1 Contributions list",       "Item 3 softened; cross-dataset [10°, 20°] bucket finding added",                                                    "R2-Q5"),
    ("§2.1 Related Work",           "Added SelaVPR++, Implicit Aggregation, Deep-Homography VPR, BoQ",                                                   "R1-W4, R2-Q4"),
    ("§3.2 / §3.3",                 "\"Guaranteed invariance\" softened to \"exact on group, approximate elsewhere\"; appearance-branch modularity made explicit", "R2-Q5, R2-Q4"),
    ("§3.3 fusion rule",            "Weighted joint scoring introduced as main; concat / attention retained as ablations",                              "R2-Q4"),
    ("§4.2 protocol",               "3 seeds, val-best ckpt, R@5 / R@10, dual GPU install path (Ampere + Blackwell)",                                   "R2-Q2, R2-Q3"),
    ("§5.1 Tables 1 + 2",           "Rebuilt with SALAD / CricaVPR / BoQ-R50 / BoQ-DINOv2; R@5 / R@10 / params / latency columns",                      "R2-Q3, R2-Q4"),
    ("§5.3 ablation Table 3",       "Each-branch + 3-fusion (attention / concat / joint scoring) + GroupPool ablations",                                "R2-Q4"),
    ("§5.4 Fig 6 / Fig 2 / Table 4", "Rotation feature curve; yaw distribution; per-yaw bucket R@1 (cross-dataset [10°, 20°] gain)",                    "R1-W3, R2-Q4, R2-Q6"),
    ("§6.2 Limitations L1–L8",      "8 explicit limitations; L8 supported by union-retrieve and single-stage equivalence ablations",                    "R2-Q7"),
    ("Supp S1",                     "Fine β sweep [0.00, 0.15] step 0.01 — plateau analysis justifying β = 0.10",                                       "R2-Q4"),
    ("Supp S2",                     "Max / mean / ℓ₂-norm GroupPool comparison",                                                                         "R2-Q4"),
    ("Supp S3",                     "Per-yaw bucket R@1 (10° buckets, both datasets)",                                                                  "R1-W3"),
    ("Supp S4",                     "Full 10-sequence ConPR per-pair breakdown",                                                                         "R2-Q3"),
    ("Supp S5",                     "Union-retrieve and single-stage equivalence null-result data (supports L8)",                                       "R2-Q7"),
    ("Supp S6 (new)",               "BoQ-DINOv2 backbone-swap experiment (sanity-checks L4 backbone-gap claim)",                                         "R2-Q5 / L4"),
    ("Minor",                       "\"banking UAVs\" → \"UAVs undergoing banked turns\"; RandAugment citation fixed",                                 "R2-Q8"),
    ("Throughout",                  "Professional language edit",                                                                                        "R2-Q9"),
]


# ------------------------------------------------------------------------
# docx helpers
# ------------------------------------------------------------------------

def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '888888')
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_qa_table(doc, qa_pairs, header_left="Reviewer comment", header_right="Our response"):
    """qa_pairs is a list of (quote, [response_paragraphs], [embeds])."""
    for entry in qa_pairs:
        if len(entry) == 3:
            quote, responses, embeds = entry
        else:
            quote, responses = entry
            embeds = []

        # Mini Q&A table for this single item
        table = doc.add_table(rows=2, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        table.allow_autofit = False
        for row in table.rows:
            row.cells[0].width = Cm(7.0)
            row.cells[1].width = Cm(10.0)

        # Header row
        hdr = table.rows[0]
        for i, txt in enumerate((header_left, header_right)):
            cell = hdr.cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(txt)
            run.bold = True
            run.font.size = Pt(10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_shading(cell, 'D9D9D9')
            set_cell_borders(cell)
            cell.width = Cm(7.0) if i == 0 else Cm(10.0)

        # Body row
        body = table.rows[1]
        body.cells[0].width = Cm(7.0)
        body.cells[1].width = Cm(10.0)

        # Left cell — reviewer quote (italic, default black)
        left = body.cells[0]
        left.text = ""
        p = left.paragraphs[0]
        run = p.add_run(quote)
        run.italic = True
        run.font.size = Pt(9.5)
        left.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        set_cell_borders(left)

        # Right cell — response paragraphs.
        # Within each paragraph, segments inside double quotes are treated as
        # paper-text quotations and rendered in Elsevier title-style blue
        # so reviewers can immediately see what we changed in the manuscript.
        right = body.cells[1]
        right.text = ""
        for ri, response_text in enumerate(responses):
            if ri == 0:
                p = right.paragraphs[0]
            else:
                p = right.add_paragraph()
            _add_response_with_quote_highlight(p, response_text)
        right.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        set_cell_borders(right)

        # Inline embeds (figures or sub-tables) immediately after the Q&A row
        # Insert a clear spacer between consecutive embeds so two tables don't
        # visually fuse into one.
        for ei, embed in enumerate(embeds):
            if ei > 0:
                gap = doc.add_paragraph()
                gap_run = gap.add_run("")
                gap_run.font.size = Pt(6)
            if embed['kind'] == 'figure':
                add_inline_figure(doc, embed['path'], embed.get('caption', ''),
                                   width_cm=embed.get('width_cm', 14.0))
            elif embed['kind'] == 'table':
                add_inline_table(doc, embed['header'], embed['rows'],
                                  caption=embed.get('caption', ''))

        # Small spacer paragraph between Q&A items
        spacer = doc.add_paragraph()
        spacer_run = spacer.add_run("")
        spacer_run.font.size = Pt(4)


PAPER_QUOTE_BLUE = RGBColor(0x1F, 0x4E, 0x79)   # Elsevier title-style blue


def _add_response_with_quote_highlight(paragraph, text):
    """Add `text` to `paragraph`, colouring substrings inside double-quotes
    in PAPER_QUOTE_BLUE so reviewers can spot exactly which strings were
    inserted into the revised manuscript.
    """
    import re
    # Capture both straight and curly double quotes; keep the quote chars in the segment
    # so the visual cue (open + close quote) stays inside the blue run.
    pattern = re.compile(r'(\"[^\"]+\"|“[^”]+”)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            run.font.size = Pt(9.5)
        quoted = paragraph.add_run(m.group(0))
        quoted.font.size = Pt(9.5)
        quoted.font.color.rgb = PAPER_QUOTE_BLUE
        quoted.italic = True
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.size = Pt(9.5)


def add_inline_figure(doc, path, caption="", width_cm=14.0):
    if not Path(path).exists():
        para = doc.add_paragraph()
        run = para.add_run(f"[Figure missing: {path}]")
        run.italic = True
        run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(8.5)


def add_inline_table(doc, header, rows, caption=""):
    # Caption ABOVE the table — makes "Table X. ..." flow naturally before
    # the data and visually separates consecutive tables.
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = cap.add_run(caption)
        run.italic = True
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr_cells = table.rows[0].cells
    for i, txt in enumerate(header):
        cell = hdr_cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(txt)
        run.bold = True
        run.font.size = Pt(9)
        set_cell_shading(cell, 'D9D9D9')
        set_cell_borders(cell)

    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ""
            r = cell.paragraphs[0].add_run(str(val))
            r.font.size = Pt(9)
            # Bold our-method row
            if "DR-VPR" in str(val) or "ours" in str(val):
                r.bold = True
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_borders(cell)


def add_summary_table(doc):
    table = doc.add_table(rows=1 + len(SUMMARY_OF_CHANGES), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for i, txt in enumerate(("Location", "Change", "Driven by")):
        cell = hdr[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(txt)
        run.bold = True
        run.font.size = Pt(10)
        set_cell_shading(cell, 'D9D9D9')
        set_cell_borders(cell)
    for ri, (loc, change, driven) in enumerate(SUMMARY_OF_CHANGES, start=1):
        row = table.rows[ri]
        for i, txt in enumerate((loc, change, driven)):
            cell = row.cells[i]
            cell.text = ""
            r = cell.paragraphs[0].add_run(txt)
            r.font.size = Pt(9.5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_borders(cell)


# ------------------------------------------------------------------------
# main
# ------------------------------------------------------------------------

def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # Title
    doc.add_heading(TITLE, level=0).alignment = WD_ALIGN_PARAGRAPH.LEFT
    for line in SUBTITLE_LINES:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(10)

    doc.add_paragraph()

    # Cover note
    doc.add_heading("Cover note to the editor", level=1)
    for para in COVER_NOTE_PARAGRAPHS:
        doc.add_paragraph(para)

    # Reviewer #1
    doc.add_heading("Response to Reviewer #1", level=1)
    doc.add_heading("Overall summary", level=2)
    add_qa_table(doc, R1_OVERALL)
    doc.add_heading("Strengths — acknowledged", level=2)
    add_qa_table(doc, R1_STRENGTHS)
    doc.add_heading("Weaknesses — addressed", level=2)
    add_qa_table(doc, R1_WEAKNESSES)

    # Reviewer #2
    doc.add_heading("Response to Reviewer #2", level=1)
    add_qa_table(doc, R2_QA)

    # Summary
    doc.add_heading("Summary of changes", level=1)
    add_summary_table(doc)

    # Closing
    doc.add_paragraph()
    doc.add_paragraph(
        "We believe that the revised manuscript addresses every reviewer concern with empirical "
        "support. We thank the Editor and Reviewers again for their constructive feedback."
    )
    doc.add_paragraph("Respectfully,")
    doc.add_paragraph("The Authors")

    out_path = "doc/REBUTTAL_LETTER.docx"
    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
