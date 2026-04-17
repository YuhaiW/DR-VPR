"""
Generate a two-column rebuttal .docx (reviewer comment | our response) by
*parsing* the authoritative text from REBUTTAL_LETTER.md. Keep the md as the
single source of truth — any edit there is reflected in the docx on rebuild.

Output: doc/REBUTTAL_LETTER.docx
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = Path(__file__).resolve().parent
MD_PATH = HERE / "REBUTTAL_LETTER.md"
OUT = HERE / "REBUTTAL_LETTER.docx"

# ---------------------------------------------------------------- parsing


def read_md() -> str:
    if not MD_PATH.is_file():
        sys.exit(f"REBUTTAL_LETTER.md not found at {MD_PATH}")
    return MD_PATH.read_text(encoding="utf-8")


SECTION_R1_SUMMARY = "## Overall summary"
SECTION_R1_STRENGTHS = "## Strengths — acknowledged"
SECTION_R1_WEAKNESSES = "## Weaknesses — addressed"
SECTION_R2 = "# Response to Reviewer #2"
SECTION_SUMMARY_CHANGES = "## Summary of changes"


def extract_cover_meta(md: str) -> dict:
    meta = {}
    for line in md.splitlines():
        m = re.match(r"^\*\*([^*]+)\*\*:\s*(.+)$", line.strip())
        if m:
            meta[m.group(1).strip()] = m.group(2).strip()
        if line.strip().startswith("---"):
            break
    return meta


def extract_cover_paragraphs(md: str) -> list[str]:
    """Everything between '## Cover note to the editor' and the next top-level header."""
    i = md.find("## Cover note to the editor")
    if i < 0:
        return []
    j = md.find("\n# Response to Reviewer #1", i)
    if j < 0:
        j = len(md)
    chunk = md[i:j]
    # drop heading line
    chunk = chunk.split("\n", 1)[1]
    # Split into paragraphs, drop leading/trailing horizontal rules
    paragraphs = []
    for para in re.split(r"\n\s*\n", chunk.strip()):
        para = para.strip()
        if not para or para == "---":
            continue
        if para.startswith("---"):
            continue
        paragraphs.append(para)
    return paragraphs


def _clean_quote(text: str) -> str:
    """Strip '>', leading/trailing quote marks, collapse whitespace."""
    # Remove leading '>' markers on each line
    lines = [re.sub(r"^\s*>\s?", "", ln) for ln in text.splitlines()]
    t = " ".join(ln.strip() for ln in lines).strip()
    # Strip bold **Qx.** prefix inside blockquote, if present
    t = re.sub(r"^\*\*[A-Za-z0-9.]+\*\*\s*", "", t)
    # Strip outer quote marks
    t = t.strip()
    if len(t) >= 2 and t[0] in "\"'“„" and t[-1] in "\"'”“":
        t = t[1:-1].strip()
    return t


_MD_BOLD = re.compile(r"\*\*(.+?)\*\*")
_MD_ITALIC = re.compile(r"(?<!\*)\*(?!\s)([^*\n]+?)\*(?!\*)")
_MD_CODE = re.compile(r"`([^`]+)`")


def _strip_markdown_inline(text: str) -> str:
    """Remove bold/italic/code markers; keep text."""
    text = _MD_BOLD.sub(r"\1", text)
    text = _MD_ITALIC.sub(r"\1", text)
    text = _MD_CODE.sub(r"\1", text)
    return text


def extract_entries(chunk: str) -> list[tuple[str, str, str]]:
    """
    Parse a chunk of markdown containing one or more:
        > **LABEL.** "quote..."    (possibly multiline)
        ...blank line...
        **Response.** response text
        ...(may contain its own blockquotes, lists, tables)...
    Returns list of (label, quote_plain, response_markdown).
    """
    entries = []
    # Split on each block that starts with `> **LABEL.` OR for R1 overall summary style
    # we accept `> "quote..."` without a label.
    # Strategy: find every "Response." occurrence and slice backwards + forwards.

    lines = chunk.splitlines()

    # Indices of Response. headers
    resp_idxs = [i for i, ln in enumerate(lines) if re.match(r"^\*\*Response\.\*\*", ln.strip())]

    for k, ri in enumerate(resp_idxs):
        # Response body: lines from ri to next Response. (or end of chunk or next section)
        next_ri = resp_idxs[k + 1] if k + 1 < len(resp_idxs) else len(lines)
        # But stop before the next top-level block: `## `, `# `, or `---`
        stop = next_ri
        for j in range(ri + 1, next_ri):
            ls = lines[j].strip()
            if ls.startswith("## ") or ls.startswith("# ") or ls.startswith("---"):
                stop = j
                break
            # Also stop if we hit a `> **Qx.**` line — that's the next quote
            if ls.startswith(">") and re.search(r"^\*\*[A-Za-z0-9.]+\*\*", ls.lstrip("> \t")):
                stop = j
                break
        response_raw = "\n".join(lines[ri:stop]).strip()

        # Quote: scan backwards from ri to find preceding blockquote that starts with `> **LABEL.`
        # or a bare `> "..."` block.
        quote_end = ri
        # find the start of the blockquote
        q_start = ri - 1
        while q_start >= 0 and not lines[q_start].strip().startswith(">"):
            q_start -= 1
        # walk backwards through contiguous '>' lines
        quote_block_end = q_start
        while q_start >= 0 and lines[q_start].strip().startswith(">"):
            q_start -= 1
        quote_block_start = q_start + 1
        if quote_block_start > quote_block_end:
            continue  # no quote found

        raw_quote_block = "\n".join(lines[quote_block_start:quote_block_end + 1])

        # Extract label from the quote block header if present
        label_match = re.search(r"^\s*>\s*\*\*([A-Za-z0-9.]+)\.\*\*", raw_quote_block, re.MULTILINE)
        label = label_match.group(1) if label_match else ""

        quote_plain = _clean_quote(raw_quote_block)

        # Strip leading `**Response.**` from response text — the docx will show it implicitly
        response = re.sub(r"^\*\*Response\.\*\*\s*", "", response_raw, count=1)
        entries.append((label, quote_plain, response))

    return entries


def extract_section_chunk(md: str, start_header: str, stop_headers: list[str]) -> str:
    i = md.find(start_header)
    if i < 0:
        return ""
    j = len(md)
    for stop in stop_headers:
        k = md.find(stop, i + len(start_header))
        if 0 <= k < j:
            j = k
    return md[i:j]


def extract_summary_rows(md: str) -> list[tuple[str, str, str]]:
    """Parse the 'Summary of changes' markdown table."""
    idx = md.find(SECTION_SUMMARY_CHANGES)
    if idx < 0:
        return []
    chunk = md[idx:]
    rows = []
    for line in chunk.splitlines():
        m = re.match(r"^\|\s*(.+?)\s*\|\s*(.+?)\s*\|\s*(.+?)\s*\|$", line)
        if m and "---" not in line and not m.group(1).strip().lower().startswith("location"):
            rows.append((m.group(1).strip(), m.group(2).strip(), m.group(3).strip()))
    return rows


def extract_closing(md: str) -> list[str]:
    # Everything after the Summary of changes table and before end
    idx = md.find(SECTION_SUMMARY_CHANGES)
    if idx < 0:
        return []
    # Find last table row end
    chunk = md[idx:]
    # The closing paragraphs come after the final '| ... |' row — find last such line
    lines = chunk.splitlines()
    last_row = -1
    for i, ln in enumerate(lines):
        if re.match(r"^\|\s*.+\s*\|\s*.+\s*\|\s*.+\s*\|$", ln):
            last_row = i
    if last_row < 0:
        return []
    tail = "\n".join(lines[last_row + 1:]).strip()
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", tail) if p.strip()]
    return paragraphs


# ---------------------------------------------------------------- rendering


def set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "BFBFBF")
        borders.append(el)
    tc_pr.append(borders)


def _add_formatted_paragraph(cell, text: str, first: bool, is_first_para: bool):
    """
    Add a paragraph of markdown-ish text to a cell. Render **bold**, *italic*,
    and `code` inline; strip block-level '>' markers; treat embedded newlines
    as soft line breaks.
    """
    text = text.rstrip()

    # Drop leading '> ' from each line (block-quote lines inside response)
    lines = [re.sub(r"^\s*>\s?", "", ln) for ln in text.splitlines()]
    text = "\n".join(lines)

    # List bullets: '- ' / '* ' at line start
    # If every non-empty line starts with '- ', render as bullet points.
    bullet_lines = [ln for ln in text.splitlines() if ln.strip()]
    is_bullet_block = bool(bullet_lines) and all(
        ln.lstrip().startswith(("- ", "* ")) for ln in bullet_lines
    )

    if is_bullet_block:
        for ln in bullet_lines:
            body = re.sub(r"^\s*[-*]\s+", "", ln)
            if is_first_para and first:
                p = cell.paragraphs[0]
                first = False
            else:
                p = cell.add_paragraph()
            p.style = "List Bullet" if "List Bullet" in [s.name for s in cell.part.document.styles] else p.style
            _emit_inline(p, body)
        return first

    # Plain paragraph with possible \n soft breaks
    if is_first_para and first:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()

    parts = text.split("\n")
    for j, part in enumerate(parts):
        if j > 0:
            p.add_run().add_break()
        _emit_inline(p, part)
    return False


def _emit_inline(p, text: str) -> None:
    """Emit a paragraph with bold / italic / code inline markup."""
    # Build tokens by successively applying regex; simplest correct approach:
    # iterate characters and maintain a small state machine. Short text here.
    i = 0
    while i < len(text):
        if text.startswith("**", i):
            end = text.find("**", i + 2)
            if end == -1:
                _emit_run(p, text[i:], bold=False, italic=False, code=False)
                break
            _emit_run(p, text[i + 2:end], bold=True, italic=False, code=False)
            i = end + 2
        elif text.startswith("`", i):
            end = text.find("`", i + 1)
            if end == -1:
                _emit_run(p, text[i:], bold=False, italic=False, code=False)
                break
            _emit_run(p, text[i + 1:end], bold=False, italic=False, code=True)
            i = end + 1
        elif text.startswith("*", i) and not text.startswith("**", i):
            # italic
            end = text.find("*", i + 1)
            if end == -1 or end == i + 1:
                _emit_run(p, text[i], bold=False, italic=False, code=False)
                i += 1
                continue
            _emit_run(p, text[i + 1:end], bold=False, italic=True, code=False)
            i = end + 1
        else:
            nxt = len(text)
            for marker in ("**", "*", "`"):
                m = text.find(marker, i + 1)
                if 0 <= m < nxt:
                    nxt = m
            _emit_run(p, text[i:nxt], bold=False, italic=False, code=False)
            i = nxt


def _emit_run(p, text: str, *, bold: bool, italic: bool, code: bool) -> None:
    if not text:
        return
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(10)
    if code:
        r.font.name = "Consolas"


def add_entry_table(doc, triples, *,
                    left_col_width_cm: float = 8.0,
                    right_col_width_cm: float = 8.5):
    """For each (label, quote, response_markdown), render bold label + 2-col table."""
    for idx, (label, quote, response) in enumerate(triples):
        lbl = doc.add_paragraph()
        lbl.paragraph_format.space_before = Pt(6 if idx > 0 else 0)
        lbl.paragraph_format.space_after = Pt(2)
        r = lbl.add_run(label or "")
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(0x1F, 0x3B, 0x70)

        table = doc.add_table(rows=2, cols=2)
        table.autofit = False
        table.allow_autofit = False
        widths = [Cm(left_col_width_cm), Cm(right_col_width_cm)]

        hdr = table.rows[0].cells
        for i, text in enumerate(("Reviewer comment", "Our response")):
            hdr[i].width = widths[i]
            set_cell_shading(hdr[i], "D9E2F3")
            set_cell_borders(hdr[i])
            p = hdr[i].paragraphs[0]
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(9.5)
            hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        row = table.rows[1].cells
        row[0].width = widths[0]
        row[1].width = widths[1]
        for cell in row:
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        # Quote cell
        p = row[0].paragraphs[0]
        qrun = p.add_run(f"“{quote}”")
        qrun.italic = True
        qrun.font.size = Pt(10)

        # Response cell — split on blank lines into paragraphs
        paragraphs = [p.strip() for p in re.split(r"\n\s*\n", response) if p.strip()]
        first = True
        for pi, para in enumerate(paragraphs):
            first = _add_formatted_paragraph(row[1], para, first, pi == 0)


def add_section_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x3B, 0x70)
    return h


# ----------------------------------------------------------------- build


def main() -> None:
    md = read_md()

    # Title (from first '# ' heading)
    m = re.match(r"#\s+(.+)", md.strip().splitlines()[0])
    title = m.group(1) if m else "Response to Reviewers"

    meta = extract_cover_meta(md)
    cover_paras = extract_cover_paragraphs(md)

    r1_summary_chunk = extract_section_chunk(md, SECTION_R1_SUMMARY, [SECTION_R1_STRENGTHS])
    r1_strengths_chunk = extract_section_chunk(md, SECTION_R1_STRENGTHS, [SECTION_R1_WEAKNESSES])
    r1_weaknesses_chunk = extract_section_chunk(md, SECTION_R1_WEAKNESSES, [SECTION_R2])
    r2_chunk = extract_section_chunk(md, SECTION_R2, [SECTION_SUMMARY_CHANGES])

    r1_sum = extract_entries(r1_summary_chunk)
    # If parser didn't attach a label for the R1 summary block, supply one
    r1_sum = [(lbl or "Summary", q, r) for (lbl, q, r) in r1_sum]

    r1_str = extract_entries(r1_strengths_chunk)
    r1_wk = extract_entries(r1_weaknesses_chunk)
    r2 = extract_entries(r2_chunk)
    summary_rows = extract_summary_rows(md)
    closing_paras = extract_closing(md)

    print(f"[parse] cover paragraphs: {len(cover_paras)}")
    print(f"[parse] R1 summary: {len(r1_sum)}, strengths: {len(r1_str)}, weaknesses: {len(r1_wk)}")
    print(f"[parse] R2 entries: {len(r2)}")
    print(f"[parse] summary rows: {len(summary_rows)}")
    print(f"[parse] closing paragraphs: {len(closing_paras)}")

    # ----------- render
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # Title
    t = doc.add_heading(title, level=0)
    for r in t.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x3B, 0x70)

    # Meta
    for key, val in meta.items():
        p = doc.add_paragraph()
        rk = p.add_run(f"{key}: ")
        rk.bold = True
        rk.font.size = Pt(10.5)
        p.add_run(val).font.size = Pt(10.5)
    doc.add_paragraph()

    # Cover note
    add_section_heading(doc, "Cover note to the editor", level=1)
    for para in cover_paras:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        _emit_inline(p, para)

    # Reviewer 1
    doc.add_page_break()
    add_section_heading(doc, "Response to Reviewer #1", level=1)
    add_section_heading(doc, "Overall summary", level=2)
    add_entry_table(doc, r1_sum)
    doc.add_paragraph()

    add_section_heading(doc, "Strengths — acknowledged", level=2)
    add_entry_table(doc, r1_str)
    doc.add_paragraph()

    add_section_heading(doc, "Weaknesses — addressed", level=2)
    add_entry_table(doc, r1_wk)

    # Reviewer 2
    doc.add_page_break()
    add_section_heading(doc, "Response to Reviewer #2", level=1)
    add_entry_table(doc, r2)

    # Summary of changes
    doc.add_page_break()
    add_section_heading(doc, "Summary of changes to the manuscript", level=1)

    sum_table = doc.add_table(rows=1, cols=3)
    sum_table.autofit = False
    sum_table.allow_autofit = False
    sum_widths = [Cm(3.2), Cm(9.5), Cm(4.0)]
    hdr = sum_table.rows[0].cells
    for i, text in enumerate(("Location", "Change", "Driven by")):
        hdr[i].width = sum_widths[i]
        set_cell_shading(hdr[i], "D9E2F3")
        set_cell_borders(hdr[i])
        p = hdr[i].paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(10)

    for loc, change, driver in summary_rows:
        row = sum_table.add_row().cells
        for i, val in enumerate((loc, change, driver)):
            row[i].width = sum_widths[i]
            set_cell_borders(row[i])
            p = row[i].paragraphs[0]
            _emit_inline(p, val)

    doc.add_paragraph()
    for para in closing_paras:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        _emit_inline(p, para)

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
